import os
from datetime import datetime


class ConversionError(Exception):
    """Raised when document conversion fails."""


def _timestamped_output_path(source_path, output_dir, target_ext):
    base_name = os.path.splitext(os.path.basename(source_path))[0]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_name = f'{base_name}_converted_{timestamp}{target_ext}'
    os.makedirs(output_dir, exist_ok=True)
    return os.path.join(output_dir, output_name)


def convert_pdf_to_docx(pdf_path, output_dir):
    """
    Convert PDF to DOCX by extracting page text and writing it into a DOCX file.

    This keeps conversion fast and robust for hackathon usage, but does not preserve
    complex original formatting.
    """
    try:
        import fitz
    except ImportError as exc:
        raise ConversionError('Missing dependency: PyMuPDF (fitz).') from exc

    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError('Missing dependency: python-docx.') from exc

    output_path = _timestamped_output_path(pdf_path, output_dir, '.docx')

    try:
        pdf = fitz.open(pdf_path)
        if len(pdf) == 0:
            raise ConversionError('PDF has no pages.')

        document = Document()
        document.add_heading('Converted from PDF', level=1)

        for page_idx in range(len(pdf)):
            page = pdf.load_page(page_idx)
            page_text = page.get_text('text').strip()

            document.add_heading(f'Page {page_idx + 1}', level=2)

            if page_text:
                for paragraph in page_text.split('\n'):
                    cleaned = paragraph.strip()
                    if cleaned:
                        document.add_paragraph(cleaned)
            else:
                document.add_paragraph('[No extractable text on this page]')

        document.save(output_path)
        pdf.close()
        return output_path
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f'PDF to DOCX conversion failed: {exc}') from exc


def _iter_docx_text_parts(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                yield ' | '.join(cells)


def convert_docx_to_pdf(docx_path, output_dir):
    """
    Convert DOCX to PDF via plain-text layout.

    The output is intentionally simple and reliable: text and table rows are rendered
    with line wrapping and page breaks.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError('Missing dependency: python-docx.') from exc

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase.pdfmetrics import stringWidth
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise ConversionError('Missing dependency: reportlab.') from exc

    output_path = _timestamped_output_path(docx_path, output_dir, '.pdf')

    try:
        doc = Document(docx_path)
        text_parts = list(_iter_docx_text_parts(doc))
        if not text_parts:
            text_parts = ['[Document has no readable text content]']

        page_width, page_height = A4
        left_margin = 50
        right_margin = 50
        top_margin = 60
        bottom_margin = 50
        line_height = 14
        font_name = 'Helvetica'
        font_size = 11

        pdf = canvas.Canvas(output_path, pagesize=A4)
        pdf.setFont(font_name, font_size)

        y = page_height - top_margin
        max_width = page_width - left_margin - right_margin

        for part in text_parts:
            wrapped_lines = []
            for raw_line in part.splitlines() or ['']:
                line = raw_line.strip()
                if not line:
                    wrapped_lines.append('')
                    continue

                if stringWidth(line, font_name, font_size) <= max_width:
                    wrapped_lines.append(line)
                    continue

                words = line.split()
                current = []
                for word in words:
                    test = (' '.join(current + [word])).strip()
                    if stringWidth(test, font_name, font_size) <= max_width:
                        current.append(word)
                    else:
                        if current:
                            wrapped_lines.append(' '.join(current))
                        current = [word]
                if current:
                    wrapped_lines.append(' '.join(current))

            for line in wrapped_lines:
                if y <= bottom_margin:
                    pdf.showPage()
                    pdf.setFont(font_name, font_size)
                    y = page_height - top_margin

                pdf.drawString(left_margin, y, line)
                y -= line_height

            # Extra spacing between paragraphs
            y -= line_height // 2

        pdf.save()
        return output_path
    except ConversionError:
        raise
    except Exception as exc:
        raise ConversionError(f'DOCX to PDF conversion failed: {exc}') from exc

def markdown_to_docx(markdown_path, output_dir):
    """
    Convert Markdown to DOCX using a simple line-by-line approach.

    This is a basic implementation that handles headings, lists, and paragraphs.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError('Missing dependency: python-docx.') from exc

    output_path = _timestamped_output_path(markdown_path, output_dir, '.docx')

    try:
        document = Document()
        with open(markdown_path, 'r', encoding='utf-8') as f:
            for line in f:
                stripped = line.strip()
                if not stripped:
                    continue
                elif stripped.startswith('#'):
                    level = len(stripped) - len(stripped.lstrip('#'))
                    document.add_heading(stripped.lstrip('#').strip(), level=level)
                elif stripped.startswith(('-', '*', '+')):
                    document.add_paragraph(stripped[1:].strip(), style='List Bullet')
                else:
                    document.add_paragraph(stripped)

        document.save(output_path)
        return output_path
    except Exception as exc:
        raise ConversionError(f'Markdown to DOCX conversion failed: {exc}') from exc